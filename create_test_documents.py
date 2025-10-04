#!/usr/bin/env python3
"""
Script to create proper test documents in correct formats
"""

import os
import json
from datetime import datetime, date
from pathlib import Path

# Create test-docs directory
os.makedirs("test-docs", exist_ok=True)

def create_pdf_content():
    """Create PDF content using reportlab"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        # Create PDF for maintenance checklist
        doc = SimpleDocTemplate("test-docs/maintenance_checklist_weekly_inspection.pdf", pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph("KOCHI METRO RAIL LIMITED", title_style))
        story.append(Paragraph("WEEKLY INSPECTION CHECKLIST", title_style))
        story.append(Spacer(1, 12))
        
        # Equipment details
        story.append(Paragraph("Equipment: Traction Motor Assembly TM-001", styles['Normal']))
        story.append(Paragraph("Date: 15-Jan-2024", styles['Normal']))
        story.append(Paragraph("Inspector: John Doe (ID: ENG-001)", styles['Normal']))
        story.append(Paragraph("Department: Engineering", styles['Normal']))
        story.append(Paragraph("Location: Aluva Depot", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Safety requirements
        story.append(Paragraph("SAFETY REQUIREMENTS:", styles['Heading2']))
        story.append(Paragraph("□ Lockout/Tagout procedures followed", styles['Normal']))
        story.append(Paragraph("□ Personal Protective Equipment (PPE) worn:", styles['Normal']))
        story.append(Paragraph("  - Safety glasses", styles['Normal']))
        story.append(Paragraph("  - Hard hat", styles['Normal']))
        story.append(Paragraph("  - Steel-toed boots", styles['Normal']))
        story.append(Paragraph("  - High-visibility vest", styles['Normal']))
        story.append(Paragraph("□ Work area properly cordoned off", styles['Normal']))
        story.append(Paragraph("□ Emergency procedures reviewed", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Inspection procedures
        story.append(Paragraph("INSPECTION PROCEDURES:", styles['Heading2']))
        story.append(Paragraph("1. VISUAL INSPECTION", styles['Heading3']))
        story.append(Paragraph("   □ Motor housing free from cracks or damage", styles['Normal']))
        story.append(Paragraph("   □ Cooling fins clean and unobstructed", styles['Normal']))
        story.append(Paragraph("   □ Electrical connections secure and dry", styles['Normal']))
        story.append(Paragraph("   □ No oil leaks or contamination", styles['Normal']))
        story.append(Paragraph("   □ Mounting bolts properly torqued", styles['Normal']))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph("2. ELECTRICAL CHECKS", styles['Heading3']))
        story.append(Paragraph("   □ Insulation resistance > 1MΩ", styles['Normal']))
        story.append(Paragraph("   □ Ground fault protection functional", styles['Normal']))
        story.append(Paragraph("   □ Terminal connections tight", styles['Normal']))
        story.append(Paragraph("   □ Cable insulation intact", styles['Normal']))
        story.append(Paragraph("   □ Control panel indicators normal", styles['Normal']))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph("3. MECHANICAL CHECKS", styles['Heading3']))
        story.append(Paragraph("   □ Bearing temperature < 80°C", styles['Normal']))
        story.append(Paragraph("   □ Vibration levels within limits", styles['Normal']))
        story.append(Paragraph("   □ Alignment within tolerance", styles['Normal']))
        story.append(Paragraph("   □ Lubrication levels adequate", styles['Normal']))
        story.append(Paragraph("   □ Drive coupling condition good", styles['Normal']))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph("4. PERFORMANCE TESTS", styles['Heading3']))
        story.append(Paragraph("   □ Start-up current within limits", styles['Normal']))
        story.append(Paragraph("   □ Speed regulation ±2%", styles['Normal']))
        story.append(Paragraph("   □ Load sharing balanced", styles['Normal']))
        story.append(Paragraph("   □ Noise levels < 85dB", styles['Normal']))
        story.append(Paragraph("   □ Efficiency > 95%", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Findings
        story.append(Paragraph("FINDINGS:", styles['Heading2']))
        story.append(Paragraph("□ No issues found", styles['Normal']))
        story.append(Paragraph("□ Minor issues noted (see remarks)", styles['Normal']))
        story.append(Paragraph("□ Major issues requiring attention", styles['Normal']))
        story.append(Paragraph("□ Equipment out of service", styles['Normal']))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph("REMARKS:", styles['Heading3']))
        story.append(Paragraph("- Bearing temperature slightly elevated at 82°C", styles['Normal']))
        story.append(Paragraph("- Recommend bearing replacement in next maintenance cycle", styles['Normal']))
        story.append(Paragraph("- All other parameters within normal limits", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("NEXT INSPECTION DUE: 22-Jan-2024", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("SIGNATURES:", styles['Heading2']))
        story.append(Paragraph("Inspector: _________________ Date: _________", styles['Normal']))
        story.append(Paragraph("Supervisor: ________________ Date: _________", styles['Normal']))
        story.append(Paragraph("Safety Officer: _____________ Date: _________", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("KMRL Form No: MAINT-001", styles['Normal']))
        story.append(Paragraph("Revision: A", styles['Normal']))
        story.append(Paragraph("Page 1 of 1", styles['Normal']))
        
        doc.build(story)
        print("✓ Created maintenance_checklist_weekly_inspection.pdf")
        
    except ImportError:
        print("❌ reportlab not installed. Installing...")
        os.system("pip install reportlab")
        return create_pdf_content()

def create_docx_content():
    """Create DOCX content using python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create maintenance schedule DOCX
        doc = Document()
        
        # Title
        title = doc.add_heading('KOCHI METRO RAIL LIMITED', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('MONTHLY MAINTENANCE SCHEDULE', 0)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('January 2024', level=1)
        
        # Department info
        doc.add_paragraph('DEPARTMENT: Engineering')
        doc.add_paragraph('PREPARED BY: Maintenance Planning Team')
        doc.add_paragraph('APPROVED BY: Chief Engineer')
        doc.add_paragraph('DATE: 01-Jan-2024')
        
        doc.add_heading('EQUIPMENT MAINTENANCE SCHEDULE:', level=1)
        
        # Week 1
        doc.add_heading('Week 1 (Jan 1-7):', level=2)
        doc.add_paragraph('• Traction Motors: TM-001, TM-002, TM-003')
        doc.add_paragraph('• Brake Systems: BS-001, BS-002')
        doc.add_paragraph('• HVAC Units: HVAC-001, HVAC-002')
        doc.add_paragraph('• Power Supply Units: PSU-001, PSU-002')
        
        # Week 2
        doc.add_heading('Week 2 (Jan 8-14):', level=2)
        doc.add_paragraph('• Signaling Equipment: SIG-001, SIG-002, SIG-003')
        doc.add_paragraph('• Communication Systems: COM-001, COM-002')
        doc.add_paragraph('• Fire Safety Systems: FIRE-001, FIRE-002')
        doc.add_paragraph('• Escalator Systems: ESC-001, ESC-002')
        
        # Week 3
        doc.add_heading('Week 3 (Jan 15-21):', level=2)
        doc.add_paragraph('• Rolling Stock: RS-001, RS-002, RS-003')
        doc.add_paragraph('• Track Maintenance: Track Section A, B, C')
        doc.add_paragraph('• Station Equipment: SE-001, SE-002, SE-003')
        doc.add_paragraph('• Security Systems: SEC-001, SEC-002')
        
        # Week 4
        doc.add_heading('Week 4 (Jan 22-28):', level=2)
        doc.add_paragraph('• Electrical Systems: ELEC-001, ELEC-002')
        doc.add_paragraph('• Water Management: WM-001, WM-002')
        doc.add_paragraph('• Waste Management: WASTE-001, WASTE-002')
        doc.add_paragraph('• Emergency Systems: EMERG-001, EMERG-002')
        
        doc.add_heading('MAINTENANCE TYPES:', level=1)
        doc.add_paragraph('1. Preventive Maintenance (PM)')
        doc.add_paragraph('2. Corrective Maintenance (CM)')
        doc.add_paragraph('3. Predictive Maintenance (PDM)')
        doc.add_paragraph('4. Emergency Maintenance (EM)')
        
        doc.add_heading('SAFETY REQUIREMENTS:', level=1)
        doc.add_paragraph('• All maintenance activities must follow KMRL Safety Protocol')
        doc.add_paragraph('• Lockout/Tagout procedures mandatory')
        doc.add_paragraph('• Personal Protective Equipment required')
        doc.add_paragraph('• Work permits for high-risk activities')
        doc.add_paragraph('• Safety briefing before each shift')
        
        doc.add_heading('RESOURCE ALLOCATION:', level=1)
        doc.add_paragraph('• Maintenance Team A: 8 technicians')
        doc.add_paragraph('• Maintenance Team B: 6 technicians')
        doc.add_paragraph('• Maintenance Team C: 4 technicians')
        doc.add_paragraph('• Supervisors: 3 per shift')
        doc.add_paragraph('• Safety Officers: 1 per shift')
        
        doc.add_heading('BUDGET ALLOCATION:', level=1)
        doc.add_paragraph('• Labor Costs: ₹2,50,000')
        doc.add_paragraph('• Spare Parts: ₹1,80,000')
        doc.add_paragraph('• Tools & Equipment: ₹45,000')
        doc.add_paragraph('• Safety Equipment: ₹25,000')
        doc.add_paragraph('• Total Budget: ₹5,00,000')
        
        doc.save('test-docs/maintenance_schedule_monthly.docx')
        print("✓ Created maintenance_schedule_monthly.docx")
        
    except ImportError:
        print("❌ python-docx not installed. Installing...")
        os.system("pip install python-docx")
        return create_docx_content()

def create_xlsx_content():
    """Create XLSX content using openpyxl"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
        from openpyxl.utils import get_column_letter
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Quarterly Budget Report"
        
        # Title
        ws['A1'] = 'KOCHI METRO RAIL LIMITED'
        ws['A1'].font = Font(size=16, bold=True)
        ws['A2'] = 'QUARTERLY BUDGET REPORT'
        ws['A2'].font = Font(size=14, bold=True)
        ws['A3'] = 'Q4 2023-24 (Jan-Mar 2024)'
        ws['A3'].font = Font(size=12, bold=True)
        
        # Department info
        ws['A5'] = 'DEPARTMENT: Finance'
        ws['A6'] = 'PREPARED BY: Finance Team'
        ws['A7'] = 'APPROVED BY: CFO'
        ws['A8'] = 'DATE: 01-Jan-2024'
        
        # Budget summary
        ws['A10'] = 'BUDGET SUMMARY:'
        ws['A10'].font = Font(bold=True)
        
        ws['A12'] = 'REVENUE:'
        ws['B12'] = 'Amount (₹)'
        ws['A13'] = 'Passenger Revenue'
        ws['B13'] = 450000000
        ws['A14'] = 'Non-Fare Revenue'
        ws['B14'] = 85000000
        ws['A15'] = 'Government Grants'
        ws['B15'] = 250000000
        ws['A16'] = 'Other Income'
        ws['B16'] = 25000000
        ws['A17'] = 'TOTAL REVENUE'
        ws['B17'] = 810000000
        ws['A17'].font = Font(bold=True)
        ws['B17'].font = Font(bold=True)
        
        ws['A19'] = 'EXPENDITURE:'
        ws['A20'] = 'Operations'
        ws['B20'] = 350000000
        ws['A21'] = 'Maintenance'
        ws['B21'] = 150000000
        ws['A22'] = 'Staff Salaries'
        ws['B22'] = 120000000
        ws['A23'] = 'Energy Costs'
        ws['B23'] = 80000000
        ws['A24'] = 'Insurance'
        ws['B24'] = 25000000
        ws['A25'] = 'Other Expenses'
        ws['B25'] = 55000000
        ws['A26'] = 'TOTAL EXPENDITURE'
        ws['B26'] = 780000000
        ws['A26'].font = Font(bold=True)
        ws['B26'].font = Font(bold=True)
        
        ws['A28'] = 'NET PROFIT'
        ws['B28'] = 30000000
        ws['A28'].font = Font(bold=True)
        ws['B28'].font = Font(bold=True)
        
        # Format columns
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 15
        
        wb.save('test-docs/financial_budget_quarterly.xlsx')
        print("✓ Created financial_budget_quarterly.xlsx")
        
    except ImportError:
        print("❌ openpyxl not installed. Installing...")
        os.system("pip install openpyxl")
        return create_xlsx_content()

def create_image_content():
    """Create PNG image using PIL"""
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # Create a simple schematic diagram
        img = Image.new('RGB', (800, 600), color='white')
        draw = ImageDraw.Draw(img)
        
        # Try to use a default font, fallback to basic if not available
        try:
            font = ImageFont.truetype("arial.ttf", 16)
        except:
            font = ImageFont.load_default()
        
        # Title
        draw.text((50, 20), "KOCHI METRO RAIL LIMITED", fill='black', font=font)
        draw.text((50, 40), "TECHNICAL SCHEMATIC DIAGRAM", fill='black', font=font)
        draw.text((50, 60), "Traction Motor Control Circuit", fill='black', font=font)
        
        # Draw schematic components
        # Power Supply Unit
        draw.rectangle([100, 150, 200, 200], outline='black', width=2)
        draw.text((120, 170), "PSU", fill='black', font=font)
        
        # Motor Control Unit
        draw.rectangle([300, 150, 400, 200], outline='black', width=2)
        draw.text((320, 170), "MCU", fill='black', font=font)
        
        # Traction Motor
        draw.rectangle([500, 150, 600, 200], outline='black', width=2)
        draw.text((520, 170), "TM", fill='black', font=font)
        
        # Brake Control Unit
        draw.rectangle([300, 300, 400, 350], outline='black', width=2)
        draw.text((320, 320), "BCU", fill='black', font=font)
        
        # Safety Systems
        draw.rectangle([100, 300, 200, 350], outline='black', width=2)
        draw.text((120, 320), "Safety", fill='black', font=font)
        
        # Communication
        draw.rectangle([500, 300, 600, 350], outline='black', width=2)
        draw.text((520, 320), "Comm", fill='black', font=font)
        
        # Draw connections
        draw.line([(200, 175), (300, 175)], fill='black', width=2)  # PSU to MCU
        draw.line([(400, 175), (500, 175)], fill='black', width=2)  # MCU to TM
        draw.line([(350, 200), (350, 300)], fill='black', width=2)  # MCU to BCU
        draw.line([(150, 200), (150, 300)], fill='black', width=2)  # PSU to Safety
        draw.line([(550, 200), (550, 300)], fill='black', width=2)  # TM to Comm
        
        # Labels
        draw.text((50, 400), "COMPONENTS:", fill='black', font=font)
        draw.text((50, 420), "PSU - Power Supply Unit", fill='black', font=font)
        draw.text((50, 440), "MCU - Motor Control Unit", fill='black', font=font)
        draw.text((50, 460), "TM - Traction Motor", fill='black', font=font)
        draw.text((50, 480), "BCU - Brake Control Unit", fill='black', font=font)
        draw.text((50, 500), "Safety - Safety Systems", fill='black', font=font)
        draw.text((50, 520), "Comm - Communication", fill='black', font=font)
        
        # Specifications
        draw.text((400, 400), "SPECIFICATIONS:", fill='black', font=font)
        draw.text((400, 420), "Power: 500kW", fill='black', font=font)
        draw.text((400, 440), "Voltage: 750V DC", fill='black', font=font)
        draw.text((400, 460), "Current: 667A", fill='black', font=font)
        draw.text((400, 480), "Speed: 1500 RPM", fill='black', font=font)
        draw.text((400, 500), "Efficiency: 95%", fill='black', font=font)
        
        img.save('test-docs/engineering_blueprint_schematic.png')
        print("✓ Created engineering_blueprint_schematic.png")
        
    except ImportError:
        print("❌ Pillow not installed. Installing...")
        os.system("pip install Pillow")
        return create_image_content()

def create_incident_report_pdf():
    """Create incident report PDF"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        doc = SimpleDocTemplate("test-docs/incident_report_signal_failure.pdf", pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1
        )
        story.append(Paragraph("KOCHI METRO RAIL LIMITED", title_style))
        story.append(Paragraph("INCIDENT REPORT", title_style))
        story.append(Spacer(1, 12))
        
        # Incident details
        story.append(Paragraph("INCIDENT ID: INC-2024-001", styles['Normal']))
        story.append(Paragraph("DATE OF INCIDENT: 15-Jan-2024", styles['Normal']))
        story.append(Paragraph("TIME: 14:30 IST", styles['Normal']))
        story.append(Paragraph("LOCATION: MG Road Station", styles['Normal']))
        story.append(Paragraph("REPORTED BY: Station Controller Rajesh Kumar (ID: SC-001)", styles['Normal']))
        story.append(Paragraph("SEVERITY: Medium", styles['Normal']))
        story.append(Paragraph("STATUS: Under Investigation", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("INCIDENT SUMMARY:", styles['Heading2']))
        story.append(Paragraph("Signal failure occurred at MG Road Station Platform 1, causing 15-minute service delay. No injuries reported. Approximately 150 passengers affected.", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("TIMELINE OF EVENTS:", styles['Heading2']))
        story.append(Paragraph("14:30 - Signal failure detected at Platform 1", styles['Normal']))
        story.append(Paragraph("14:32 - Station Controller notified Control Room", styles['Normal']))
        story.append(Paragraph("14:35 - Manual signal operation initiated", styles['Normal']))
        story.append(Paragraph("14:40 - Emergency response team dispatched", styles['Normal']))
        story.append(Paragraph("14:45 - Train service restored with manual operation", styles['Normal']))
        story.append(Paragraph("15:00 - Normal automated signaling restored", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("ROOT CAUSE ANALYSIS:", styles['Heading2']))
        story.append(Paragraph("Preliminary investigation indicates electrical fault in signal control system. Suspected causes:", styles['Normal']))
        story.append(Paragraph("1. Power supply fluctuation", styles['Normal']))
        story.append(Paragraph("2. Control module malfunction", styles['Normal']))
        story.append(Paragraph("3. Communication link failure", styles['Normal']))
        story.append(Paragraph("4. Environmental factors (humidity/temperature)", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("AFFECTED SYSTEMS:", styles['Heading2']))
        story.append(Paragraph("• Platform signaling system", styles['Normal']))
        story.append(Paragraph("• Train control system", styles['Normal']))
        story.append(Paragraph("• Passenger information system", styles['Normal']))
        story.append(Paragraph("• Emergency communication system", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("IMPACT ASSESSMENT:", styles['Heading2']))
        story.append(Paragraph("• Service delay: 15 minutes", styles['Normal']))
        story.append(Paragraph("• Passengers affected: 150", styles['Normal']))
        story.append(Paragraph("• Revenue impact: ₹15,000 (estimated)", styles['Normal']))
        story.append(Paragraph("• Reputation impact: Minimal", styles['Normal']))
        story.append(Paragraph("• Safety impact: None", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("CORRECTIVE ACTIONS TAKEN:", styles['Heading2']))
        story.append(Paragraph("1. Immediate manual operation of signals", styles['Normal']))
        story.append(Paragraph("2. Emergency response team activation", styles['Normal']))
        story.append(Paragraph("3. Passenger information updates", styles['Normal']))
        story.append(Paragraph("4. Alternative routing implemented", styles['Normal']))
        story.append(Paragraph("5. Technical team dispatched for repair", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("SIGNATURES:", styles['Heading2']))
        story.append(Paragraph("Station Controller: _______________ Date: _________", styles['Normal']))
        story.append(Paragraph("Control Room Supervisor: _________ Date: _________", styles['Normal']))
        story.append(Paragraph("Safety Officer: _________________ Date: _________", styles['Normal']))
        story.append(Paragraph("Technical Head: _________________ Date: _________", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("KMRL Form No: INC-001", styles['Normal']))
        story.append(Paragraph("Revision: C", styles['Normal']))
        story.append(Paragraph("Classification: Internal Use", styles['Normal']))
        
        doc.build(story)
        print("✓ Created incident_report_signal_failure.pdf")
        
    except ImportError:
        print("❌ reportlab not installed. Installing...")
        os.system("pip install reportlab")
        return create_incident_report_pdf()

def main():
    """Create all test documents in proper formats"""
    print("Creating proper test documents...")
    
    # Remove old text files
    for file in Path("test-docs").glob("*"):
        if file.suffix in ['.pdf', '.docx', '.xlsx', '.png', '.dwg']:
            file.unlink()
    
    # Create proper documents
    create_pdf_content()
    create_incident_report_pdf()
    create_docx_content()
    create_xlsx_content()
    create_image_content()
    
    print("\n✅ All test documents created successfully!")
    print("📁 Documents are now in proper binary formats:")
    print("   - PDF files are actual PDFs")
    print("   - DOCX files are actual Word documents")
    print("   - XLSX files are actual Excel spreadsheets")
    print("   - PNG files are actual images")

if __name__ == "__main__":
    main()
